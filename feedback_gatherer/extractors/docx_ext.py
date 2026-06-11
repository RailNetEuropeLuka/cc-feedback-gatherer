"""Extract .docx feedback in four shapes:
  1. Table responses   - a table like  Paragraph | Considerations | Proposals  (Trenitalia).
  2. "same as FTE"     - identical template endorsing the FTE response; flagged, not re-parsed.
  3. Review markup     - Word comment balloons and tracked changes (insertions/deletions).
                         python-docx ignores these, yet for reviewed copies of the
                         Guidelines they ARE the feedback. Parsed from the raw XML.
  4. Prose             - section-numbered paragraphs, split on numbered headers.

If the file is an RNE-published copy of the Guidelines itself (letterhead markers
from config), its body text is the *guideline* text - not feedback - so body-derived
items are suppressed and only markup (3) is kept.
"""
from __future__ import annotations

import io
import re
import zipfile
from xml.etree import ElementTree as ET

import docx

from .base import ParsedSource, RawItem, topic_items, outline_items

EMAIL_RE = re.compile(r"[\w.+-]+@[\w.-]+\.\w+")
_W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
_HDR = re.compile(r"^\s*([1-3]\.\d(?:\.\d+)?)\s+\S")

# header cell keywords that mark the three logical columns of a feedback table
_PARA_KEYS = ("paragraph", "section", "chapter", "article", "ref")
_CONS_KEYS = ("consideration", "comment", "remark", "input", "feedback", "observation")
_PROP_KEYS = ("proposal", "suggestion", "amendment", "change", "wording")


def _email(text: str) -> str | None:
    m = EMAIL_RE.search(text or "")
    return m.group(0) if m else None


def _classify_table(header_cells: list[str]) -> dict | None:
    """If the header row looks like a feedback table, return a column->role map."""
    roles = {}
    for i, h in enumerate(header_cells):
        h = (h or "").lower()
        if any(k in h for k in _PARA_KEYS):
            roles.setdefault("section", i)
        elif any(k in h for k in _CONS_KEYS):
            roles.setdefault("considerations", i)
        elif any(k in h for k in _PROP_KEYS):
            roles.setdefault("proposal", i)
    return roles if "considerations" in roles else None


# --------------------------------------------------------------- review markup
def _para_text(p) -> str:
    return "".join(t.text or "" for t in p.iter(f"{_W}t")).strip()


def _markup_items(data: bytes) -> list[RawItem]:
    """Word comments + tracked changes, with the guideline section each one sits in.

    Walks word/document.xml paragraph-by-paragraph in document order, tracking the
    last numbered section header seen, so every comment anchor and tracked change
    can be tagged with its section.
    """
    z = zipfile.ZipFile(io.BytesIO(data))
    names = z.namelist()

    comments_by_id: dict[str, tuple[str, str]] = {}   # id -> (author, text)
    if "word/comments.xml" in names:
        croot = ET.fromstring(z.read("word/comments.xml"))
        for c in croot.findall(f"{_W}comment"):
            cid = c.get(f"{_W}id", "")
            author = c.get(f"{_W}author", "")
            text = " ".join(_para_text(p) for p in c.findall(f"{_W}p")).strip()
            if text:
                comments_by_id[cid] = (author, text)

    items: list[RawItem] = []
    root = ET.fromstring(z.read("word/document.xml"))
    section = ""
    for p in root.iter(f"{_W}p"):
        ptext = _para_text(p)
        if _HDR.match(ptext):
            section = ptext[:90]

        # comment anchors in this paragraph
        for ref in p.iter(f"{_W}commentReference"):
            cid = ref.get(f"{_W}id", "")
            if cid in comments_by_id:
                author, text = comments_by_id.pop(cid)
                items.append(RawItem(
                    section_raw=section or "general",
                    considerations=text,
                    raw_text=(f"Comment by {author or 'unknown'}"
                              + (f' on: "{ptext[:250]}"' if ptext else "") + f" -> {text}"),
                    confidence="medium" if section else "low"))

        # tracked changes in this paragraph
        ins = " ".join(t.text or "" for el in p.findall(f".//{_W}ins")
                       for t in el.iter(f"{_W}t")).strip()
        dele = " ".join(t.text or "" for el in p.findall(f".//{_W}del")
                        for t in el.iter(f"{_W}delText")).strip()
        if ins or dele:
            change = []
            if ins:
                change.append(f'insert "{ins[:400]}"')
            if dele:
                change.append(f'delete "{dele[:400]}"')
            items.append(RawItem(
                section_raw=section or "general",
                considerations=f"Tracked change in: {ptext[:250]}" if ptext
                               else "Tracked change",
                proposal=" / ".join(change),
                raw_text=f"Tracked change ({' / '.join(change)})"
                         + (f' in paragraph: "{ptext[:250]}"' if ptext else ""),
                confidence="medium" if section else "low"))

    # comments whose anchor was not found still carry feedback
    for author, text in comments_by_id.values():
        items.append(RawItem(section_raw="general", considerations=text,
                             raw_text=f"Comment by {author or 'unknown'} -> {text}",
                             confidence="low"))
    return items


def extract(data: bytes, filename: str, cfg: dict) -> list[ParsedSource]:
    doc = docx.Document(io.BytesIO(data))
    paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    full_text = "\n".join(paras)

    ps = ParsedSource(source_format="docx")
    ps.email_hint = _email(full_text)
    # company / representative hint: first 1-3 short lines before the body
    if paras:
        ps.company_hint = paras[0][:120]
        for p in paras[:4]:
            if " - " in p or "–" in p:
                ps.company_hint = p[:120]
                break

    # ---- (2) "same as FTE" endorsement -----------------------------------
    markers = [m.lower() for m in cfg.get("same_as_fte_filename_markers", ["same as fte"])]
    if any(m in filename.lower() for m in markers):
        ps.is_endorsement = True
        ps.notes.append("Endorses the FTE response (\"same as FTE\").")
        ps.full_text = full_text
        return [ps]

    # ---- (3) review markup: comments + tracked changes --------------------
    try:
        markup = _markup_items(data)
    except Exception as exc:
        markup = []
        ps.notes.append(f"Could not parse review markup: {exc!r}")

    head = full_text[:4000].lower()
    if any(m.lower() in head for m in cfg.get("guidelines_doc_markers", [])):
        # marked-up copy of the Guidelines: the body is guideline text, not feedback
        ps.company_hint = None
        ps.notes.append("Marked-up copy of the Guidelines: body text suppressed, "
                        f"{len(markup)} comment(s)/tracked change(s) extracted as feedback.")
        ps.items = markup
        ps.full_text = full_text
        if not markup:
            ps.notes.append("No comments or tracked changes found in the Guidelines copy.")
        return [ps]

    # ---- (1) feedback table ----------------------------------------------
    parts = []
    for t in doc.tables:
        if not t.rows:
            continue
        header = [c.text.strip() for c in t.rows[0].cells]
        roles = _classify_table(header)
        if not roles:
            continue
        sc, cc, pc = roles.get("section"), roles["considerations"], roles.get("proposal")
        last_section = "general remarks"
        for r in t.rows[1:]:
            cells = [c.text.strip() for c in r.cells]
            cons = cells[cc] if cc < len(cells) else ""
            if not cons:
                continue
            section = cells[sc] if (sc is not None and sc < len(cells) and cells[sc]) else last_section
            last_section = section
            prop = cells[pc] if (pc is not None and pc < len(cells)) else None
            ps.items.append(RawItem(section_raw=section, considerations=cons,
                                    proposal=(prop or None),
                                    raw_text=" | ".join(c for c in cells if c)))
            parts.append(f"### {section}\n{cons}" + (f"\nProposal: {prop}" if prop else ""))

    # ---- (4) prose fallback (only when neither tables nor markup matched) --
    if not ps.items and not markup:
        _split_prose(full_text, ps)
        if not ps.items:
            outline = outline_items(full_text, cfg["sections"])
            if outline:
                ps.notes.append(f"Split on the document's own outline: {len(outline)} "
                                "item(s) anchored to guideline chapters via the headings.")
                ps.items.extend(outline)
        if not ps.items:
            topics = topic_items(full_text)
            if topics:
                ps.notes.append(f"No numbered sections; split on {len(topics)} "
                                "topic heading(s) instead (canonical section unknown).")
                ps.items.extend(topics)

    if markup:
        ps.notes.append(f"{len(markup)} comment(s)/tracked change(s) extracted "
                        "in addition to the document text.")
        ps.items.extend(markup)

    ps.full_text = "\n\n".join(parts) if parts else full_text
    if not ps.items:
        ps.notes.append("No structured feedback detected; stored as one general item.")
        ps.items.append(RawItem(section_raw="general", considerations=full_text,
                                raw_text=full_text, confidence="low"))
    return [ps]


def _split_prose(text: str, ps: ParsedSource):
    """Split section-numbered prose into items keyed by the numbered headers."""
    from taxonomy import Taxonomy  # local import to avoid cycle at module load
    matches = list(Taxonomy.HEADER_RE.finditer(text))
    if not matches:
        return
    for i, m in enumerate(matches):
        start = m.end()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
        body = text[start:end].strip()
        if body:
            ps.items.append(RawItem(section_raw=f"{m.group(1)} {m.group(2)}".strip(),
                                    considerations=body, raw_text=m.group(0) + "\n" + body,
                                    confidence="medium"))
